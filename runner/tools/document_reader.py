"""
Document Reader Tool

Legge e converte in testo file di vari formati:
PDF, DOCX, XLSX/XLS, CSV, TXT, MD, JSON, YAML
"""
import csv
import json
import io
from pathlib import Path
from typing import Dict, Any, List, Optional
from loguru import logger

from .base import Tool

# Mappa estensioni → parser
SUPPORTED_FORMATS = {
    "text": [".txt", ".md", ".rst", ".log", ".ini", ".cfg", ".toml", ".env"],
    "code": [".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".css", ".sh", ".yaml", ".yml", ".json"],
    "pdf": [".pdf"],
    "word": [".docx", ".doc"],
    "excel": [".xlsx", ".xls", ".ods"],
    "csv": [".csv"],
    "rtf": [".rtf"],
}

MAX_FILE_SIZE_MB = 50


class DocumentReaderTool(Tool):
    """Tool per leggere documenti di vari formati"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__(
            name="document_reader",
            description=(
                "Read and extract text content from files of various formats: "
                "PDF, Word (DOCX), Excel (XLSX), CSV, plain text, code files, JSON, YAML. "
                "Returns the full text content ready to be analyzed."
            ),
            parameters={
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute or ~ path to the file to read"
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Max characters to return (default: 100000). Use 0 for no limit."
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "For Excel files: sheet name to read (default: first sheet)"
                    }
                },
                "required": ["path"]
            }
        )
        self.config = config
        cfg_perms = config.get("permissions", {}).get("filesystem", {})
        self.max_size_mb = cfg_perms.get("max_file_size_mb", MAX_FILE_SIZE_MB)

    def execute(
        self,
        path: str,
        max_chars: int = 100_000,
        sheet_name: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        target = Path(path).expanduser().resolve()

        if not target.exists():
            return {"success": False, "error": f"File not found: {target}"}
        if not target.is_file():
            return {"success": False, "error": f"Not a file: {target}"}

        size_mb = target.stat().st_size / (1024 * 1024)
        if size_mb > self.max_size_mb:
            return {"success": False, "error": f"File too large: {size_mb:.1f} MB (limit {self.max_size_mb} MB)"}

        suffix = target.suffix.lower()
        logger.info(f"Reading document: {target} ({suffix}, {size_mb:.2f} MB)")

        try:
            if suffix in SUPPORTED_FORMATS["pdf"]:
                text, meta = self._read_pdf(target)
            elif suffix in SUPPORTED_FORMATS["word"]:
                text, meta = self._read_word(target)
            elif suffix in SUPPORTED_FORMATS["excel"]:
                text, meta = self._read_excel(target, sheet_name)
            elif suffix in SUPPORTED_FORMATS["csv"]:
                text, meta = self._read_csv(target)
            elif suffix in SUPPORTED_FORMATS["text"] + SUPPORTED_FORMATS["code"]:
                text, meta = self._read_text(target)
            else:
                # Try as plain text anyway
                text, meta = self._read_text(target)
                meta["format"] = "unknown (treated as text)"

            if max_chars and len(text) > max_chars:
                text = text[:max_chars] + f"\n\n[... truncated at {max_chars} chars ...]"

            return {
                "success": True,
                "path": str(target),
                "format": meta.get("format", suffix),
                "size_mb": round(size_mb, 3),
                "char_count": len(text),
                "metadata": meta,
                "content": text,
            }

        except Exception as e:
            logger.error(f"Error reading {target}: {e}")
            return {"success": False, "error": str(e), "path": str(target)}

    # ── Parsers ──────────────────────────────────────────────────────────────

    def _read_pdf(self, path: Path):
        try:
            import pdfplumber
            pages_text = []
            metadata = {}
            with pdfplumber.open(str(path)) as pdf:
                metadata = {
                    "format": "pdf",
                    "pages": len(pdf.pages),
                    "info": dict(pdf.metadata) if pdf.metadata else {}
                }
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text() or ""
                    if t.strip():
                        pages_text.append(f"--- Page {i} ---\n{t}")
            return "\n\n".join(pages_text), metadata
        except ImportError:
            # Fallback to pypdf
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                pages_text = []
                for i, page in enumerate(reader.pages, 1):
                    t = page.extract_text() or ""
                    if t.strip():
                        pages_text.append(f"--- Page {i} ---\n{t}")
                return "\n\n".join(pages_text), {
                    "format": "pdf",
                    "pages": len(reader.pages)
                }
            except ImportError:
                raise ImportError("Install pdfplumber or pypdf: pip install pdfplumber")

    def _read_word(self, path: Path):
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also read tables
            tables_text = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables_text.append("\n".join(rows))
            text = "\n".join(paragraphs)
            if tables_text:
                text += "\n\n=== Tables ===\n" + "\n\n---\n".join(tables_text)
            return text, {
                "format": "docx",
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables)
            }
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

    def _read_excel(self, path: Path, sheet_name: Optional[str] = None):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            sheets = wb.sheetnames

            if sheet_name and sheet_name in sheets:
                target_sheets = [sheet_name]
            else:
                target_sheets = sheets

            all_text = []
            for sname in target_sheets:
                ws = wb[sname]
                rows_text = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows_text.append(" | ".join(cells))
                if rows_text:
                    all_text.append(f"=== Sheet: {sname} ===\n" + "\n".join(rows_text))

            return "\n\n".join(all_text), {
                "format": "xlsx",
                "sheets": sheets,
                "sheets_read": target_sheets
            }
        except ImportError:
            raise ImportError("Install openpyxl: pip install openpyxl")

    def _read_csv(self, path: Path):
        rows = []
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows), {
            "format": "csv",
            "rows": len(rows)
        }

    def _read_text(self, path: Path):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return text, {"format": path.suffix.lstrip(".") or "txt"}


def get_file_format(path: str) -> str:
    """Utility: ritorna il formato del file"""
    suffix = Path(path).suffix.lower()
    for fmt, exts in SUPPORTED_FORMATS.items():
        if suffix in exts:
            return fmt
    return "unknown"


def is_readable(path: str) -> bool:
    """Utility: controlla se il file è leggibile"""
    suffix = Path(path).suffix.lower()
    all_exts = [ext for exts in SUPPORTED_FORMATS.values() for ext in exts]
    return suffix in all_exts
