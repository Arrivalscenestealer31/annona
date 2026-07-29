"""
Tests for DocumentReaderTool.

Covers: TXT, CSV, JSON, Markdown, XLSX (with openpyxl),
        DOCX (with python-docx), PDF (with pdfplumber),
        missing files, oversized files, truncation, unknown extensions.
"""

import json
from unittest.mock import MagicMock, patch

import pytest

from runner.tools.document_reader import DocumentReaderTool, get_file_format, is_readable

# ── Helpers ───────────────────────────────────────────────────────────────────


def make_tool(config=None):
    cfg = config or {"permissions": {"filesystem": {"max_file_size_mb": 50}}}
    return DocumentReaderTool(cfg)


# ── Plain text / code formats ─────────────────────────────────────────────────


class TestTextFormats:
    def test_read_txt(self, tmp_path):
        f = tmp_path / "hello.txt"
        f.write_text("Hello, world!\nLine 2.\n")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "Hello, world!" in result["content"]
        assert result["format"] == "txt"

    def test_read_markdown(self, tmp_path):
        f = tmp_path / "README.md"
        f.write_text("# Title\n\nSome **bold** text.\n")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "# Title" in result["content"]
        assert result["format"] == "md"

    def test_read_json(self, tmp_path):
        data = {"key": "value", "number": 42, "list": [1, 2, 3]}
        f = tmp_path / "data.json"
        f.write_text(json.dumps(data, indent=2))
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert '"key"' in result["content"]
        assert result["char_count"] > 0

    def test_read_python_script(self, tmp_path):
        f = tmp_path / "main.py"
        f.write_text("def hello():\n    return 'world'\n")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "def hello" in result["content"]

    def test_read_shell_script(self, tmp_path):
        f = tmp_path / "deploy.sh"
        f.write_text("#!/bin/bash\necho 'deploying'\n")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "deploying" in result["content"]

    def test_truncation(self, tmp_path):
        f = tmp_path / "big.txt"
        f.write_text("x" * 10_000)
        result = make_tool().execute(str(f), max_chars=100)
        assert result["success"] is True
        assert "truncated" in result["content"]
        assert len(result["content"]) < 200  # 100 chars + truncation message

    def test_no_truncation_when_zero(self, tmp_path):
        f = tmp_path / "big.txt"
        content = "y" * 5_000
        f.write_text(content)
        result = make_tool().execute(str(f), max_chars=0)
        assert result["success"] is True
        assert result["content"] == content

    def test_unicode_content(self, tmp_path):
        f = tmp_path / "unicode.txt"
        f.write_text("Ciao 🌍 café résumé 日本語\n", encoding="utf-8")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "café" in result["content"]

    def test_unknown_extension_treated_as_text(self, tmp_path):
        f = tmp_path / "data.xyz"
        f.write_text("some raw content")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "some raw content" in result["content"]


# ── CSV ───────────────────────────────────────────────────────────────────────


class TestCSV:
    def test_basic_csv(self, tmp_path):
        f = tmp_path / "data.csv"
        f.write_text("Name,Age,City\nAlice,30,Rome\nBob,25,Milan\n")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "Alice" in result["content"]
        assert "Milan" in result["content"]
        assert result["metadata"]["rows"] == 3  # header + 2 data rows

    def test_csv_with_quotes(self, tmp_path):
        f = tmp_path / "quoted.csv"
        f.write_text('id,description\n1,"Product with, comma"\n2,"Another ""quoted"" item"\n')
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert "Product with, comma" in result["content"]

    def test_empty_csv(self, tmp_path):
        f = tmp_path / "empty.csv"
        f.write_text("")
        result = make_tool().execute(str(f))
        assert result["success"] is True
        assert result["content"] == ""


# ── Excel (openpyxl) ──────────────────────────────────────────────────────────


class TestExcel:
    @pytest.fixture
    def xlsx_file(self, tmp_path):
        pytest.importorskip("openpyxl")
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["Product", "Q1", "Q2", "Q3"])
        ws.append(["Widget A", 100, 200, 300])
        ws.append(["Widget B", 150, 250, 350])
        ws2 = wb.create_sheet("Team")
        ws2.append(["Name", "Role"])
        ws2.append(["Alice", "Engineer"])
        ws2.append(["Bob", "Designer"])
        path = tmp_path / "workbook.xlsx"
        wb.save(str(path))
        return path

    def test_reads_all_sheets(self, xlsx_file):
        result = make_tool().execute(str(xlsx_file))
        assert result["success"] is True
        assert "Sales" in result["content"]
        assert "Widget A" in result["content"]
        assert "Team" in result["content"]
        assert "Alice" in result["content"]

    def test_reads_specific_sheet(self, xlsx_file):
        result = make_tool().execute(str(xlsx_file), sheet_name="Team")
        assert result["success"] is True
        assert "Alice" in result["content"]
        assert "Widget A" not in result["content"]  # Sales sheet not included

    def test_metadata_has_sheet_names(self, xlsx_file):
        result = make_tool().execute(str(xlsx_file))
        assert "Sales" in result["metadata"]["sheets"]
        assert "Team" in result["metadata"]["sheets"]


# ── DOCX (python-docx) ────────────────────────────────────────────────────────


class TestDocx:
    @pytest.fixture
    def docx_file(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Annual Report 2024", level=1)
        doc.add_paragraph("Executive Summary")
        doc.add_paragraph("Revenue grew by 25% year-over-year.")
        # Add a table
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Quarter"
        table.cell(0, 1).text = "Revenue"
        table.cell(1, 0).text = "Q1"
        table.cell(1, 1).text = "1,200,000"
        table.cell(2, 0).text = "Q2"
        table.cell(2, 1).text = "1,500,000"
        path = tmp_path / "report.docx"
        doc.save(str(path))
        return path

    def test_reads_paragraphs(self, docx_file):
        result = make_tool().execute(str(docx_file))
        assert result["success"] is True
        assert "Annual Report" in result["content"]
        assert "Revenue grew" in result["content"]

    def test_reads_table(self, docx_file):
        result = make_tool().execute(str(docx_file))
        assert "1,200,000" in result["content"]

    def test_metadata_paragraph_count(self, docx_file):
        result = make_tool().execute(str(docx_file))
        assert result["metadata"]["tables"] == 1
        assert result["metadata"]["paragraphs"] >= 3


# ── PDF (pdfplumber) ──────────────────────────────────────────────────────────


class TestPDF:
    @pytest.fixture
    def pdf_file(self, tmp_path):
        """Creates a minimal PDF using fpdf2 or reportlab if available,
        otherwise mocks pdfplumber."""
        try:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            pdf.cell(200, 10, txt="Annual Financial Report", ln=True)
            pdf.cell(200, 10, txt="Total Revenue: 5,000,000", ln=True)
            pdf.cell(200, 10, txt="Net Profit: 1,200,000", ln=True)
            path = tmp_path / "report.pdf"
            pdf.output(str(path))
            return path
        except ImportError:
            # Fallback: create a dummy file and mock pdfplumber
            path = tmp_path / "report.pdf"
            path.write_bytes(b"%PDF-1.4 fake")
            return path

    def test_read_pdf_with_pdfplumber_mock(self, tmp_path):
        """Test PDF reading by mocking pdfplumber (no real PDF needed)."""
        path = tmp_path / "mock.pdf"
        path.write_bytes(b"%PDF-1.4 minimal")

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page 1 content\nRevenue: 5,000,000"

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.metadata = {"Author": "Akaion", "Title": "Report"}
        mock_pdf.__enter__ = lambda s: mock_pdf
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = make_tool().execute(str(path))

        assert result["success"] is True
        assert "Revenue: 5,000,000" in result["content"]
        assert result["metadata"]["pages"] == 1
        assert result["metadata"]["info"]["Author"] == "Akaion"

    def test_read_pdf_multipage_mock(self, tmp_path):
        """Multiple pages are concatenated with page markers."""
        path = tmp_path / "multipage.pdf"
        path.write_bytes(b"%PDF-1.4")

        pages = []
        for i in range(3):
            p = MagicMock()
            p.extract_text.return_value = f"Content of page {i + 1}"
            pages.append(p)

        mock_pdf = MagicMock()
        mock_pdf.pages = pages
        mock_pdf.metadata = {}
        mock_pdf.__enter__ = lambda s: mock_pdf
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = make_tool().execute(str(path))

        assert result["success"] is True
        assert "Page 1" in result["content"]
        assert "Page 2" in result["content"]
        assert "Page 3" in result["content"]

    def test_read_pdf_empty_pages(self, tmp_path):
        """Pages with no text are silently skipped."""
        path = tmp_path / "empty_pages.pdf"
        path.write_bytes(b"%PDF-1.4")

        p1 = MagicMock()
        p1.extract_text.return_value = "   "  # whitespace only
        p2 = MagicMock()
        p2.extract_text.return_value = "Actual content"

        mock_pdf = MagicMock()
        mock_pdf.pages = [p1, p2]
        mock_pdf.metadata = {}
        mock_pdf.__enter__ = lambda s: mock_pdf
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = make_tool().execute(str(path))

        assert "Actual content" in result["content"]
        assert "Page 1" not in result["content"]  # empty page skipped


# ── Error cases ───────────────────────────────────────────────────────────────


class TestErrorCases:
    def test_file_not_found(self):
        result = make_tool().execute("/nonexistent/path/file.txt")
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_path_is_directory(self, tmp_path):
        result = make_tool().execute(str(tmp_path))
        assert result["success"] is False
        assert "not a file" in result["error"].lower()

    def test_file_too_large(self, tmp_path):
        f = tmp_path / "huge.txt"
        f.write_bytes(b"x")
        cfg = {"permissions": {"filesystem": {"max_file_size_mb": 0}}}  # 0 MB limit
        result = DocumentReaderTool(cfg).execute(str(f))
        assert result["success"] is False
        assert "large" in result["error"].lower()

    def test_tilde_expansion(self, tmp_path, monkeypatch):
        monkeypatch.setenv("HOME", str(tmp_path))
        f = tmp_path / "home_file.txt"
        f.write_text("home content")
        result = make_tool().execute("~/home_file.txt")
        assert result["success"] is True
        assert "home content" in result["content"]


# ── Utility functions ─────────────────────────────────────────────────────────


class TestUtilities:
    def test_get_file_format_pdf(self):
        assert get_file_format("report.pdf") == "pdf"

    def test_get_file_format_docx(self):
        assert get_file_format("doc.docx") == "word"

    def test_get_file_format_xlsx(self):
        assert get_file_format("sheet.xlsx") == "excel"

    def test_get_file_format_csv(self):
        assert get_file_format("data.csv") == "csv"

    def test_get_file_format_py(self):
        assert get_file_format("script.py") == "code"

    def test_get_file_format_unknown(self):
        assert get_file_format("binary.bin") == "unknown"

    def test_is_readable_known(self):
        for ext in [".txt", ".pdf", ".docx", ".xlsx", ".csv", ".py", ".md", ".json"]:
            assert is_readable(f"file{ext}") is True

    def test_is_readable_unknown(self):
        assert is_readable("file.bin") is False
        assert is_readable("file.exe") is False
